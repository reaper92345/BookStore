
from docx import Document
import os

doc_path = "ONLINE BOOKSTORE MANAGEMENT SYSTEM_UPDATED.docx"
final_output_path = "ONLINE BOOKSTORE MANAGEMENT SYSTEM_FINAL.docx"

def get_placeholder_replacements():
    return {
        "[Your Name]": "Student",
        "[Supervisor Name]": "Project Supervisor",
        "[Insert Date]": "December 2025",
        "[Swastik College (TU)]": "Swastik College (TU)"
    }

def get_tech_stack_replacements():
    return {
        "Apache": "Uvicorn",
        "FastAPI (Python)-Apache": "FastAPI (Python)-Uvicorn",
        "image: php:8.2-apache": "image: python:3.9-slim",
        "container_name: bookstore_php": "container_name: bookstore_backend",
        "- ./php:/var/www/html/": "- ./backend:/app",
        "mysqli": "SQLAlchemy",
        "db.php": "database.py",
        "Location: dashboard.php": "Redirect to Dashboard",
        "if(isset($_POST['login']))": "def login_user(db: Session, user: schemas.UserLogin):",
        "$_POST['email']": "user.email",
        "$_POST['password']": "user.password",
        "$_SESSION['user']": "access_token",
        "let cart = JSON.parse(localStorage.getItem(\"cart\")) || [];": "const cart = useSelector((state: RootState) => state.cart.items);"
    }

def polish_document():
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return

    doc = Document(doc_path)
    
    # Replacement map
    replacements = {**get_placeholder_replacements(), **get_tech_stack_replacements()}

    # Update Paragraphs
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    # Handle PHP code blocks specifically if they are in multiple paragraphs
    # (Simple replacement might be enough if they are in single lines)
    
    # Check for <?php blocks
    for i in range(len(doc.paragraphs)):
        if "<?php" in doc.paragraphs[i].text:
            doc.paragraphs[i].text = "from fastapi import FastAPI, Depends, HTTPException\nfrom sqlalchemy.orm import Session"
        if "include 'db.php';" in doc.paragraphs[i].text:
            doc.paragraphs[i].text = "from .database import get_db, engine"

    # Also check inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            para.text = para.text.replace(old, new)

    doc.save(final_output_path)
    print(f"Successfully polished document: {final_output_path}")

if __name__ == "__main__":
    polish_document()
